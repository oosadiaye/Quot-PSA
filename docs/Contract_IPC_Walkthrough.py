"""
Generate the Contract → Mobilization → Milestone → IPC → Retention
walkthrough as a Word document for business stakeholders.

Run:
    python docs/Contract_IPC_Walkthrough.py

Outputs:
    Contract_IPC_Walkthrough.docx in the user's Documents folder.
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches


# ── Theme ──────────────────────────────────────────────────────────────
BRAND_NAVY = RGBColor(0x1c, 0x20, 0x6d)        # Quot brand navy
BRAND_GREEN = RGBColor(0x39, 0xcd, 0x9a)
ACCENT_INDIGO = RGBColor(0x4f, 0x46, 0xe5)
TEXT_DARK = RGBColor(0x0f, 0x17, 0x2a)
TEXT_MUTED = RGBColor(0x64, 0x74, 0x8b)
TABLE_HEADER_BG = '1c206d'
TABLE_TOTAL_BG = '0f172a'
TABLE_ROW_ALT_BG = 'f8fafc'


def shade_cell(cell, fill_hex: str) -> None:
    """Apply a background fill colour to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.color.rgb = BRAND_NAVY
    run.font.name = 'Calibri'
    run.font.size = Pt({1: 18, 2: 14, 3: 12}.get(level, 11))
    run.bold = True


def add_para(doc: Document, text: str, *, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT_DARK
    if italic:
        run.italic = True


def add_callout(doc: Document, label: str, text: str) -> None:
    """Inline callout — bold label + body text in a single coloured block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    label_run = p.add_run(f'{label}  ')
    label_run.bold = True
    label_run.font.color.rgb = ACCENT_INDIGO
    label_run.font.size = Pt(10)
    body_run = p.add_run(text)
    body_run.font.size = Pt(10)
    body_run.font.color.rgb = TEXT_DARK


def add_journal_block(doc: Document, lines: list[tuple[str, str, str, str]]) -> None:
    """Render a journal entry as a fixed-width-style block.

    Each tuple = (DR/CR marker, gl_code, account_name, amount_string)
    """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for marker, gl, name, amount in lines:
        run = p.add_run(f'  {marker}   {gl:<12}{name:<40}{amount:>20}\n')
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = TEXT_DARK


def add_step_table(doc: Document, headers: list[str], rows: list[list[str]],
                   col_widths_cm: list[float] | None = None) -> None:
    """Branded table — navy header, alternating row fills."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # Header
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.font.size = Pt(9)
        shade_cell(cell, TABLE_HEADER_BG)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(9.5)
            run.font.color.rgb = TEXT_DARK
            # Mark TOTAL rows by checking the first cell's content
            if r_idx == len(rows) - 1 and val == 'TOTAL':
                run.bold = True
            elif r_idx % 2 == 1:
                shade_cell(cell, TABLE_ROW_ALT_BG)

    # Column widths
    if col_widths_cm:
        for c_idx, w in enumerate(col_widths_cm):
            for cell in table.columns[c_idx].cells:
                cell.width = Cm(w)


def add_section_break(doc: Document) -> None:
    """Visual separator between major sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4f46e5')
    pBdr.append(bottom)
    pPr.append(pBdr)


# ──────────────────────────────────────────────────────────────────────
# Build the document
# ──────────────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Default body font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── Cover ─────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover.add_run('CONTRACT IPC ACCOUNTING WALKTHROUGH')
    cover_run.bold = True
    cover_run.font.size = Pt(24)
    cover_run.font.color.rgb = BRAND_NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run('When journal entries post during the contract lifecycle')
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = TEXT_MUTED

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_run = org.add_run('Office of Accountant General — Delta State')
    org_run.font.size = Pt(11)
    org_run.font.color.rgb = TEXT_MUTED

    add_section_break(doc)

    # ── Setup ─────────────────────────────────────────────────────────
    add_heading(doc, 'The Setup', level=2)
    add_para(
        doc,
        'A contract is awarded for ₦2,500,000 with the following terms. We follow '
        'the entire flow from activation through final retention release.',
    )
    add_step_table(
        doc,
        headers=['Item', 'Amount'],
        rows=[
            ['Contract sum', '₦2,500,000'],
            ['Mobilization advance (40%)', '₦1,000,000'],
            ['Total retention pool (5% × contract)', '₦125,000'],
            ['Available for milestones', '₦1,375,000'],
            ['Milestone 1', '₦500,000'],
            ['Milestone 2', '₦500,000'],
            ['Milestone 3', '₦375,000'],
        ],
        col_widths_cm=[8, 4.5],
    )

    add_callout(
        doc,
        'KEY PRINCIPLE.',
        'Three distinct events trigger journal posting in this lifecycle: '
        '(1) Mobilization PV posts — advance disbursed, recognised as a receivable. '
        '(2) IPC approved — accrual journal hits the books with retention as a liability '
        'and mobilization recovery netted off. (3) PV for the IPC posts — cash leaves the '
        'TSA, AP cleared, WHT recognised. There is no journal at contract activation, '
        'milestone definition, or IPC submit — those are control records only. '
        'The contractor "earning" the money triggers the accrual; the cash leaving '
        'triggers the cash entry.',
    )

    # ── Step 1 ────────────────────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'Step 1 — Contract Activated  (no journal yet)', level=2)
    add_para(doc, 'Activate contract DSG/WORKS/2026/001.')
    add_para(
        doc,
        'GL impact: NONE. IPSAS does not recognise expense for a signed contract '
        'until work is performed.',
    )
    add_para(
        doc,
        'What happens internally: the ContractBalance row is materialised — ceiling = '
        '2,500,000; pending_voucher = 0; certified = 0. The appropriation is marked '
        'ENCUMBERED for ₦2.5M (memo only, no GL).',
        italic=True, size=10,
    )

    # ── Step 2 ────────────────────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'Step 2 — Mobilization Advance Issued + Paid', level=2)
    add_para(doc, 'This is the FIRST journal posting on the contract.')
    add_para(
        doc,
        'Action: "+ Issue Mobilization" creates a PENDING MobilizationPayment of '
        '₦1,000,000. Treasury raises a Payment Voucher; the journal posts on payment.',
    )
    add_heading(doc, 'Journal entry on PV post', level=3)
    add_journal_block(doc, [
        ('DR', '31xxxxxx', 'Mobilization Advance Receivable', '1,000,000.00'),
        ('CR', '31010100', 'TSA / Cash', '1,000,000.00'),
    ])
    add_callout(
        doc,
        'WHY A RECEIVABLE.',
        'The contractor now owes us ₦1M of future work. The asset stays on the balance '
        'sheet until IPCs progressively recover it. Mobilization is not an expense — '
        'it is an advance against work to be done.',
    )
    add_para(
        doc,
        'ContractBalance after: mobilization_paid = 1,000,000; mobilization_recovered = 0.',
        italic=True, size=10,
    )

    # ── Step 3 ────────────────────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'Step 3 — Milestone 1 (₦500,000) → IPC Approved', level=2)
    add_para(
        doc,
        'Milestone creation and approval have NO GL impact. The journal posts when '
        'IPC #1 is APPROVED.',
    )
    add_heading(doc, 'IPC #1 calculations', level=3)
    add_step_table(
        doc,
        headers=['Item', 'Amount'],
        rows=[
            ['Gross certified', '₦500,000'],
            ['Less: Retention deduction (5%)', '₦25,000'],
            ['Less: Mobilization recovery (40%)', '₦200,000'],
            ['Net to contractor', '₦275,000'],
        ],
        col_widths_cm=[8, 4.5],
    )
    add_heading(doc, 'Journal posted on Approve (IPSAS accrual)', level=3)
    add_journal_block(doc, [
        ('DR', '22020306', 'Expense (Printing/Works)', '500,000.00'),
        ('CR', '31xxxxxx', 'Mobilization Advance Receivable', '200,000.00'),
        ('CR', '41xxxxxx', 'Retention Held (Liability)', '25,000.00'),
        ('CR', '21010xxx', 'Accounts Payable — Vendor', '275,000.00'),
    ])
    add_para(
        doc,
        'Effect: Expense recognised (work has been done). Mobilization receivable '
        'shrinks by ₦200,000 — now ₦800,000. Retention liability grows by ₦25,000 — we '
        'now owe this back to the contractor at completion. AP grows by ₦275,000 — what '
        'we still owe the contractor in cash.',
        italic=True, size=10,
    )
    add_heading(doc, 'Journal posted on PV pay (cash settlement)', level=3)
    add_journal_block(doc, [
        ('DR', '21010xxx', 'Accounts Payable', '275,000.00'),
        ('CR', '41030107', 'WHT Payable (5%)', '13,750.00'),
        ('CR', '31010100', 'TSA / Cash', '261,250.00'),
    ])
    add_para(
        doc,
        'WHT = 5% × 275,000 (deferred from IPC time per cash-basis FIRS rule). '
        'AP cleared. Net cash to contractor: ₦261,250.',
        italic=True, size=10,
    )

    # ── Step 4 ────────────────────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'Step 4 — Milestone 2 (₦500,000) → IPC Approved', level=2)
    add_para(doc, 'Identical pattern to Step 3.')
    add_heading(doc, 'IPC #2 accrual journal', level=3)
    add_journal_block(doc, [
        ('DR', '22020306', 'Expense', '500,000.00'),
        ('CR', '31xxxxxx', 'Mobilization Advance', '200,000.00'),
        ('CR', '41xxxxxx', 'Retention Held', '25,000.00'),
        ('CR', '21010xxx', 'Accounts Payable — Vendor', '275,000.00'),
    ])
    add_heading(doc, 'PV pay journal (same shape as IPC #1 pay)', level=3)
    add_journal_block(doc, [
        ('DR', '21010xxx', 'Accounts Payable', '275,000.00'),
        ('CR', '41030107', 'WHT Payable', '13,750.00'),
        ('CR', '31010100', 'TSA / Cash', '261,250.00'),
    ])
    add_para(
        doc,
        'Cumulative position now: Mobilization recovered = ₦400,000 (₦600,000 still owed '
        'by contractor). Retention held = ₦50,000. Cash paid to contractor so far = '
        '₦1,000,000 advance + ₦261,250 + ₦261,250 = ₦1,522,500.',
        italic=True, size=10,
    )

    # ── Step 5 ────────────────────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'Step 5 — Milestone 3 (₦375,000) — Final IPC', level=2)
    add_para(
        doc,
        'A policy decision matters here: how does the system handle the mobilization '
        'that has not been fully recovered yet?',
    )
    add_callout(
        doc,
        'POLICY NOTE.',
        'At a flat 40% recovery rate per IPC, total recovery would be ₦550,000 — leaving '
        '₦450,000 of mobilization unrecovered when the contract ends. That is not how '
        'a Delta State contract should close. The correct policy is to recover the '
        'mobilization at a rate that fully amortises by the final IPC. Using '
        'gross × (mobilization ÷ Σ milestones) = gross × 72.73% achieves this. '
        'The figures below assume the corrected policy.',
    )
    add_heading(doc, 'IPC #3 calculations (full-recovery policy)', level=3)
    add_step_table(
        doc,
        headers=['Item', 'Amount'],
        rows=[
            ['Gross certified', '₦375,000'],
            ['Less: Retention (5%)', '₦18,750'],
            ['Less: Mobilization recovery (72.73%)', '₦272,727'],
            ['Net to contractor', '₦83,523'],
        ],
        col_widths_cm=[8, 4.5],
    )
    add_heading(doc, 'IPC #3 accrual journal', level=3)
    add_journal_block(doc, [
        ('DR', '22020306', 'Expense', '375,000.00'),
        ('CR', '31xxxxxx', 'Mobilization Advance', '272,727.00'),
        ('CR', '41xxxxxx', 'Retention Held', '18,750.00'),
        ('CR', '21010xxx', 'Accounts Payable — Vendor', '83,523.00'),
    ])
    add_heading(doc, 'PV pay journal', level=3)
    add_journal_block(doc, [
        ('DR', '21010xxx', 'Accounts Payable', '83,523.00'),
        ('CR', '41030107', 'WHT Payable (5%)', '4,176.00'),
        ('CR', '31010100', 'TSA / Cash', '79,347.00'),
    ])

    # ── Step 6 ────────────────────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'Step 6 — Practical Completion (Retention Release 1)', level=2)
    add_para(
        doc,
        'Contract status flips to PRACTICAL_COMPLETION. The "Release Retention" '
        'button on the contract detail page becomes active.',
    )
    add_step_table(
        doc,
        headers=['Item', 'Amount'],
        rows=[
            ['Total retention held', '₦68,750'],
            ['  IPC #1 retention', '₦25,000'],
            ['  IPC #2 retention', '₦25,000'],
            ['  IPC #3 retention', '₦18,750'],
            ['50% release at Practical Completion', '₦34,375'],
        ],
        col_widths_cm=[8, 4.5],
    )
    add_heading(doc, 'Journal posted on PV pay (retention release)', level=3)
    add_journal_block(doc, [
        ('DR', '41xxxxxx', 'Retention Held (Liability)', '34,375.00'),
        ('CR', '31010100', 'TSA / Cash', '34,375.00'),
    ])
    add_para(
        doc,
        'Effect: liability reduces; cash leaves the TSA to the contractor.',
        italic=True, size=10,
    )

    # ── Step 7 ────────────────────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'Step 7 — Final Completion (Retention Release 2)', level=2)
    add_para(
        doc,
        'After the defects-liability period elapses, contract status flips to '
        'FINAL_COMPLETION and the remaining 50% of retention is released.',
    )
    add_heading(doc, 'Journal posted on PV pay', level=3)
    add_journal_block(doc, [
        ('DR', '41xxxxxx', 'Retention Held', '34,375.00'),
        ('CR', '31010100', 'TSA / Cash', '34,375.00'),
    ])
    add_para(
        doc,
        'Retention liability now zero. Contract is closeable.',
        italic=True, size=10,
    )

    # ── Cumulative ────────────────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'Cumulative GL Picture', level=2)
    add_para(
        doc,
        'The full ledger position across the entire contract lifecycle. '
        'Negative values shown in parentheses follow accounting convention.',
    )
    add_step_table(
        doc,
        headers=['Stage', 'DR Expense', 'DR Mob Adv', 'DR AP', 'DR Retention', 'CR Cash', 'CR WHT'],
        rows=[
            ['Mob paid', '', '1,000,000', '', '', '(1,000,000)', ''],
            ['IPC#1 accrual', '500,000', '(200,000)', '', '', '', ''],
            ['IPC#1 pay', '', '', '275,000', '', '(261,250)', '(13,750)'],
            ['IPC#2 accrual', '500,000', '(200,000)', '', '', '', ''],
            ['IPC#2 pay', '', '', '275,000', '', '(261,250)', '(13,750)'],
            ['IPC#3 accrual', '375,000', '(272,727)', '', '', '', ''],
            ['IPC#3 pay', '', '', '83,523', '', '(79,347)', '(4,176)'],
            ['Retention release 1', '', '', '', '34,375', '(34,375)', ''],
            ['Retention release 2', '', '', '', '34,375', '(34,375)', ''],
            ['TOTAL', '1,375,000', '327,273', '633,523', '68,750', '(1,670,597)', '(31,676)'],
        ],
        col_widths_cm=[3.5, 2.4, 2.4, 2.2, 2.4, 2.4, 2.0],
    )
    add_para(
        doc,
        'Note: figures assume the corrected mobilization recovery rate. With the '
        'default pro-rata recovery, the mobilization receivable would still hold '
        '₦450,000 at contract end — that is the gap the auditor will flag.',
        italic=True, size=10,
    )

    # ── System vs. Practice ───────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'System Behaviour vs. Public Sector Practice', level=2)
    add_step_table(
        doc,
        headers=['Concern', 'Current code', 'Correct PFM practice'],
        rows=[
            ['GL impact on contract activation', 'None ✓', 'None (commitment only) ✓'],
            ['GL impact on milestone creation', 'None ✓', 'None (planning record) ✓'],
            ['GL impact on milestone approve', 'None ✓', 'None (engineer cert only) ✓'],
            ['GL impact on IPC approve', 'Posts accrual journal ✓', 'Same ✓'],
            ['GL impact on PV post', 'DR AP / CR WHT / CR Cash ✓', 'Same ✓'],
            ['Mobilization recovery formula',
             'mob_rate × gross (40% × gross), capped at remaining — leaves orphan receivable',
             'mob_amount × gross ÷ (contract − retention) — fully amortises'],
            ['Retention release at Practical', 'Button + PENDING release record ✓', 'Same ✓'],
            ['Retention release at Final', 'Same flow for remaining 50% ✓', 'Same ✓'],
        ],
        col_widths_cm=[4.5, 5.5, 5.5],
    )

    # ── Recommendation ────────────────────────────────────────────────
    add_section_break(doc)
    add_heading(doc, 'Recommendation', level=2)
    add_para(
        doc,
        'The journal-posting model is correct end-to-end. The single behavioural gap '
        'that auditors will flag is the mobilization recovery rate. A small change to '
        'MobilizationService.compute_recovery() to use the formula '
        'mob_amount × this_gross ÷ (contract_ceiling − retention_total) would auto-'
        'amortise the advance completely by the final IPC, matching Delta State '
        'practice — without changing any of the journal-posting logic above.',
    )

    # ── Closing ───────────────────────────────────────────────────────
    add_section_break(doc)
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing_run = closing.add_run(
        'Document generated from the Quot PSE platform — '
        'Contract IPC walkthrough.'
    )
    closing_run.italic = True
    closing_run.font.size = Pt(9)
    closing_run.font.color.rgb = TEXT_MUTED

    return doc


def main() -> Path:
    out = Path(os.path.expanduser('~/Documents/Contract_IPC_Walkthrough.docx'))
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = build()
    doc.save(out)
    return out


if __name__ == '__main__':
    p = main()
    print(f'Generated: {p}')
