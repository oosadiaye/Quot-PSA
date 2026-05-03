"""
Generate: Contracts & Milestone Payment — Process Flow (Word document).

Produces docs/Contracts_Milestone_Payment_Process.docx with the full, authoritative
step-by-step lifecycle as implemented by quot_pse, including:
  - the distinction between GRN (procurement of goods) and IPC (works/services)
  - the full contract state machine (DRAFT -> CLOSED)
  - the IPC state machine (DRAFT -> PAID)
  - mobilization, retention, variations
  - a textual process-flow diagram suitable for a Word document

Run:
    python scripts/generate_contract_process_doc.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ──────────────────────────────────────────────────────────────────────
# Styling helpers
# ──────────────────────────────────────────────────────────────────────

BRAND_NAVY = RGBColor(0x0B, 0x2E, 0x4F)
BRAND_ACCENT = RGBColor(0x1F, 0x6F, 0xB5)
GREY = RGBColor(0x6B, 0x72, 0x80)


def set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_NAVY


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False,
             size: int = 11, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color


def add_bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        for r in p.runs:
            r.font.size = Pt(11)


def add_numbered(doc: Document, items: list[str]) -> None:
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        for r in p.runs:
            r.font.size = Pt(11)


def add_table(doc: Document, header: list[str], rows: list[list[str]],
              *, widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    # header
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        set_cell_shading(hdr[i], "0B2E4F")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # body
    for ri, row in enumerate(rows, start=1):
        cells = table.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(10)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths_cm:
        for row in table.rows:
            for ci, w in enumerate(widths_cm):
                row.cells[ci].width = Cm(w)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "E8F1FA")
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.color.rgb = BRAND_NAVY
    r1.font.size = Pt(11)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(10)


def add_flow_step(doc: Document, n: int, title: str, actor: str, system: str) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    widths = [Cm(1.2), Cm(5.5), Cm(9.5)]

    num_cell, title_cell, sys_cell = table.rows[0].cells
    set_cell_shading(num_cell, "1F6FB5")
    set_cell_shading(title_cell, "F4F8FC")
    set_cell_shading(sys_cell, "FFFFFF")

    for c, w in zip(table.rows[0].cells, widths):
        c.width = w

    num_cell.text = ""
    p = num_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(n))
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    title_cell.text = ""
    pt = title_cell.paragraphs[0]
    rt = pt.add_run(title)
    rt.bold = True
    rt.font.size = Pt(11)
    rt.font.color.rgb = BRAND_NAVY
    pa = title_cell.add_paragraph()
    ra = pa.add_run(f"Actor: {actor}")
    ra.italic = True
    ra.font.size = Pt(9)
    ra.font.color.rgb = GREY

    sys_cell.text = ""
    ps = sys_cell.paragraphs[0]
    rs = ps.add_run(system)
    rs.font.size = Pt(10)

    # spacer
    doc.add_paragraph()


# ──────────────────────────────────────────────────────────────────────
# Content
# ──────────────────────────────────────────────────────────────────────

def build_document(out_path: Path) -> None:
    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Quot PSE — Contracts & Milestone Payment")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = BRAND_NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("End-to-End Process Flow — Start to Finish & Payment")
    rs.italic = True
    rs.font.size = Pt(13)
    rs.font.color.rgb = GREY

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run("Public Sector ERP · Nigeria IFMIS · Delta State Pilot\n"
                       "Module: contracts  |  Related: accounting, procurement, workflow")
    rm.font.size = Pt(10)
    rm.font.color.rgb = GREY

    doc.add_paragraph()

    # ── Section 1: Does Contract / IPC go through GRN? ───────────────
    add_heading(doc, "1. Does a Contract / IPC go through a GRN?", level=1)
    add_para(doc,
             "Short answer: No. A Goods Received Note (GRN) is a procurement "
             "artefact for goods-only purchase orders. Capital-works, consultancy, "
             "and services contracts that operate on physical progress (roads, "
             "buildings, maintenance, engineering studies) DO NOT pass through a "
             "GRN. Their evidence of delivery is the Engineer's Interim Payment "
             "Certificate (IPC).")

    add_para(doc, "Two parallel P2P rails exist in Quot PSE:", bold=True)
    add_table(
        doc,
        header=["Rail", "Trigger", "Evidence of delivery", "Matching", "Payment doc"],
        rows=[
            ["Procurement (goods)",
             "Purchase Requisition → Purchase Order",
             "Goods Received Note (GRN)",
             "3-way match: PO + GRN + Invoice",
             "Payment Voucher"],
            ["Contracts (works/services)",
             "Contract award → Activate",
             "Interim Payment Certificate (IPC)",
             "3-way match: Contract + IPC + Voucher-gross",
             "Payment Voucher"],
        ],
        widths_cm=[3.0, 3.5, 3.5, 3.8, 2.5],
    )

    add_callout(
        doc,
        "Key distinction",
        "The IPC is the works-equivalent of a GRN + invoice rolled into one. "
        "The certifying Engineer attests that work has physically been executed to "
        "the certified percentage, and the IPC itself becomes the claim on the "
        "treasury. Quot PSE enforces this separation via ALLOWED_IPC_TRANSITIONS "
        "in contracts/models/payment.py."
    )

    # ── Section 2: Contract State Machine ────────────────────────────
    doc.add_page_break()
    add_heading(doc, "2. Contract Lifecycle (State Machine)", level=1)
    add_para(doc,
             "Every contract advances through seven states. Transitions are "
             "enforced by ALLOWED_CONTRACT_TRANSITIONS; any attempt to skip a "
             "state is rejected with a clear error code.")

    add_table(
        doc,
        header=["#", "State", "Meaning", "Unlocks"],
        rows=[
            ["1", "DRAFT", "Awarded, not yet operational.",
             "Edit; Activate"],
            ["2", "ACTIVATED", "Contract number assigned, balance ledger materialised.",
             "Mobilization advance; IPCs; Variations"],
            ["3", "IN_PROGRESS", "Physical execution underway.",
             "Periodic IPCs; Variations; Measurement Books"],
            ["4", "PRACTICAL_COMPLETION", "Engineer has issued Practical Completion Certificate.",
             "Release of first half of retention"],
            ["5", "DEFECTS_LIABILITY", "Defects Liability Period (DLP) running.",
             "Defect rectification; second half of retention still held"],
            ["6", "FINAL_COMPLETION", "End of DLP; Final Completion Certificate issued.",
             "Release of final retention"],
            ["7", "CLOSED", "Terminal. All payments reconciled.",
             "Audit only"],
        ],
        widths_cm=[0.8, 3.0, 6.8, 4.6],
    )

    # ── Section 3: IPC State Machine ─────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "3. IPC Lifecycle (State Machine)", level=1)
    add_para(doc,
             "The Interim Payment Certificate is the payable unit for works "
             "contracts. Each IPC moves linearly through seven states and can "
             "never skip a stage. Rejections kick the IPC back to DRAFT for rework.")

    add_table(
        doc,
        header=["#", "IPC State", "Actor", "Action"],
        rows=[
            ["1", "DRAFT", "Site Engineer / Contractor rep.",
             "Prepare quantities, rates, measurement references."],
            ["2", "SUBMITTED", "Site Engineer",
             "Submit to Certifier. Integrity hash is frozen."],
            ["3", "CERTIFIER_REVIEWED", "Project Engineer / Certifier",
             "Cross-check against Measurement Book. Cannot be the same user as the submitter (SoD)."],
            ["4", "APPROVED", "Procurement Officer / Approver",
             "Financial approval. Cannot be the certifier (SoD)."],
            ["5", "VOUCHER_RAISED", "Accounts Payable",
             "Raise Payment Voucher in Treasury module. Three-way-match validated: voucher_gross == IPC.net_payable."],
            ["6", "PAID", "Treasury / Cashier",
             "Mark paid. VAT and WHT recorded at this point (FIRS cash-basis). Accrual journal posted."],
            ["—", "REJECTED", "Any reviewer",
             "Returns to DRAFT for rework with rejection_reason captured."],
        ],
        widths_cm=[0.8, 3.5, 4.5, 6.5],
    )

    add_callout(
        doc,
        "Segregation of Duties (SoD)",
        "Enforced via contracts.services.sod.actor_can_bypass_sod. Default rule: "
        "Submitter ≠ Certifier ≠ Approver. Superusers and users with an explicit "
        "bypass permission may override, but the override is audit-logged."
    )

    # ── Section 4: Full Process Flow ─────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "4. Full Process Flow — Start to Finish", level=1)
    add_para(doc,
             "The numbered steps below are the real sequence executed by the "
             "services layer and observed end-to-end by scripts/e2e_contracts_workflow.py.",
             italic=True, color=GREY)

    steps: list[tuple[str, str, str]] = [
        ("Create Contract (DRAFT)",
         "Procurement Officer",
         "POST /api/v1/contracts/contracts/. Required: title, reference, contract_type, "
         "procurement_method, vendor, MDA, NCoA code, appropriation, fiscal_year, "
         "original_sum, mobilization_rate (≤15%), retention_rate (typically 10%), "
         "signed_date, contract_start_date, contract_end_date, defects_liability_period_days."),

        ("Define Milestone Schedule",
         "Project Engineer",
         "Attach MilestoneSchedule rows: milestone_number, description, scheduled_value, "
         "percentage_weight, target_date. Sum of scheduled_value must ≤ original_sum. "
         "Milestones drive valuation of IPC line items."),

        ("Activate Contract (ACTIVATED)",
         "MDA Head + Finance",
         "POST /contracts/{id}/activate/. Generates official contract_number, "
         "materialises ContractBalance ledger (mob_advance, retention_held, "
         "certified_to_date, paid_to_date = 0), opens IPC and variation endpoints."),

        ("(Optional) Mobilization Advance",
         "Accounts Payable",
         "MobilizationService.raise_voucher → mark_paid. Advance = mobilization_rate × "
         "original_sum. Recovered pro-rata from each subsequent IPC's gross via "
         "mobilization_recovery_this_cert."),

        ("Contract moves to IN_PROGRESS",
         "Project Engineer",
         "Contract.transition_to(IN_PROGRESS). Triggered when first IPC is submitted "
         "or via explicit transition call."),

        ("Prepare IPC (DRAFT)",
         "Site Engineer",
         "Capture cumulative_work_done_to_date, previous_certified, variation_claims, "
         "ld_deduction. System computes: this_certificate_gross, retention_deduction_this_cert, "
         "mobilization_recovery_this_cert, net_payable."),

        ("Submit IPC (DRAFT → SUBMITTED)",
         "Site Engineer",
         "IPCService.submit_ipc. Integrity hash of line items is frozen; "
         "any subsequent edit must re-enter via REJECTED → DRAFT."),

        ("Engineer Certifies (SUBMITTED → CERTIFIER_REVIEWED)",
         "Certifier (different user)",
         "IPCService.certify. SoD guard blocks submitter from certifying. "
         "Measurement Book reference can be attached for defensibility."),

        ("Financial Approval (CERTIFIER_REVIEWED → APPROVED)",
         "Approver (different user)",
         "IPCService.approve. Control checks: cumulative_certified ≤ contract_ceiling; "
         "ceiling = original_sum + Σ(approved variations)."),

        ("Raise Payment Voucher (APPROVED → VOUCHER_RAISED)",
         "Accounts Payable",
         "Create PaymentVoucherGov (voucher_number, payment_type, tsa_account, "
         "gross_amount, wht_amount, net_amount, narration). "
         "IPCService.raise_voucher performs the 3-way match: "
         "abs(voucher_gross - ipc.net_payable) ≤ COHERENCE_TOLERANCE, else ThreeWayMatchError."),

        ("Mark Paid (VOUCHER_RAISED → PAID)",
         "Treasury / Cashier",
         "IPCService.mark_paid(payment_date, vat_amount, wht_amount). "
         "Posts accrual journal; updates ContractBalance.cumulative_gross_paid; "
         "records FIRS tax liability (cash-basis). Retention is held; not paid to vendor."),

        ("(If any) Variation Orders",
         "Engineer + Approver Tier",
         "ContractVariation auto-computes approval_tier: LOCAL ≤15%, BOARD ≤25%, "
         "BPP_REQUIRED >25% of original_sum. Only APPROVED variations raise the ceiling."),

        ("Practical Completion",
         "Engineer",
         "Issue CompletionCertificate (PRACTICAL). Contract → PRACTICAL_COMPLETION. "
         "RetentionService releases the first half: raise_voucher → mark_paid."),

        ("Defects Liability Period",
         "MDA + Vendor",
         "Contract → DEFECTS_LIABILITY. Duration = contract.defects_liability_period_days. "
         "Defect rectifications are zero-value IPCs or variations."),

        ("Final Completion",
         "Engineer",
         "End of DLP + no outstanding defects. Final CompletionCertificate issued. "
         "Contract → FINAL_COMPLETION."),

        ("Release Final Retention",
         "Accounts Payable + Treasury",
         "RetentionService.raise_voucher → mark_paid on the remaining retention pool. "
         "ContractBalance.retention_held → 0."),

        ("Close Contract",
         "Procurement Officer",
         "POST /contracts/{id}/close/. Control checks: cumulative_gross_certified == "
         "cumulative_gross_paid + retention_released; no open IPCs; no pending variations. "
         "Contract → CLOSED (terminal)."),
    ]
    for i, (t, a, s) in enumerate(steps, start=1):
        add_flow_step(doc, i, t, a, s)

    # ── Section 5: Textual Process-Flow Diagram ──────────────────────
    doc.add_page_break()
    add_heading(doc, "5. Process-Flow Diagram (Text)", level=1)
    add_para(doc,
             "ASCII representation suitable for pasting into presentations or "
             "rendering to graphviz. Dashed edges are optional; solid edges are "
             "mandatory state transitions.",
             italic=True, color=GREY)

    diagram = (
        "CONTRACT LIFECYCLE\n"
        "------------------\n"
        "   [Award]\n"
        "      |\n"
        "      v\n"
        "  (DRAFT) --activate--> (ACTIVATED) --first IPC--> (IN_PROGRESS)\n"
        "                             |                            |\n"
        "                             |                            v\n"
        "                             |               (PRACTICAL_COMPLETION)\n"
        "                             |                            |\n"
        "                             |                            v\n"
        "                             |                 (DEFECTS_LIABILITY)\n"
        "                             |                            |\n"
        "                             |                            v\n"
        "                             |                  (FINAL_COMPLETION)\n"
        "                             |                            |\n"
        "                             +---- retention & close -----+\n"
        "                                          |\n"
        "                                          v\n"
        "                                      (CLOSED)\n"
        "\n"
        "IPC LIFECYCLE (per payment period)\n"
        "----------------------------------\n"
        "  (DRAFT) --submit--> (SUBMITTED) --certify--> (CERTIFIER_REVIEWED)\n"
        "      ^                   |                        |\n"
        "      |                reject                   reject\n"
        "      |                   v                        v\n"
        "      +--- (REJECTED) <--+----- approve -----> (APPROVED)\n"
        "                                                   |\n"
        "                                            raise voucher\n"
        "                                                   v\n"
        "                                         (VOUCHER_RAISED)\n"
        "                                                   |\n"
        "                                                mark paid\n"
        "                                                   v\n"
        "                                                (PAID)\n"
        "\n"
        "WORKS vs GOODS (Payment Evidence)\n"
        "---------------------------------\n"
        "  GOODS (procurement):    PR -> PO -> GRN -> Invoice -> 3-way-match -> PV -> Pay\n"
        "  WORKS (contracts):      Contract -> Activate -> IPC (cert'd) ------> PV -> Pay\n"
        "                                                   ^\n"
        "                                                   +-- IPC replaces GRN+Invoice\n"
    )
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # ── Section 6: Financial Controls Summary ────────────────────────
    doc.add_page_break()
    add_heading(doc, "6. Financial Controls & Computations", level=1)
    add_table(
        doc,
        header=["Control", "Where enforced", "Formula / Rule"],
        rows=[
            ["Ceiling never breached",
             "IPCService.approve, ContractVariation.save",
             "Σ(certified) ≤ original_sum + Σ(approved variations)."],
            ["Retention held on every IPC",
             "IPCService._compute_deductions",
             "retention_deduction_this_cert = this_certificate_gross × retention_rate%."],
            ["Mobilization recovery",
             "MobilizationService.reconcile_payment_status",
             "recovery_i = advance × (gross_i / original_sum) until fully recovered."],
            ["Three-way match",
             "IPCService.raise_voucher",
             "abs(voucher_gross − ipc.net_payable) ≤ COHERENCE_TOLERANCE else ThreeWayMatchError."],
            ["Variation tiered approval",
             "ContractVariation.save (auto tier)",
             "LOCAL ≤15%, BOARD (15–25%], BPP_REQUIRED >25% of original_sum."],
            ["Tax on payment",
             "IPCService.mark_paid",
             "VAT (7.5%) and WHT (typically 5% works) recorded at payment date (FIRS cash-basis)."],
            ["SoD on IPC",
             "contracts.services.sod",
             "Submitter ≠ Certifier ≠ Approver unless bypass permission granted."],
            ["Integrity hash",
             "InterimPaymentCertificate._compute_hash",
             "SHA-256 of line-items frozen at submit; tamper detection on re-read."],
        ],
        widths_cm=[4.0, 4.5, 6.5],
    )

    # ── Section 7: Worked Example ────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "7. Worked Example — ₦10 M Road Contract", level=1)
    add_para(doc,
             "Following is the actual sequence exercised by the end-to-end harness "
             "(scripts/e2e_contracts_workflow.py) on contract E2E-001.",
             italic=True, color=GREY)

    add_table(
        doc,
        header=["Step", "Amount", "Running Certified", "Running Paid", "Retention Held"],
        rows=[
            ["Contract awarded (original_sum)", "₦10,000,000", "—", "—", "—"],
            ["Contract activated", "—", "₦0", "₦0", "₦0"],
            ["IPC #1 gross", "₦2,500,000", "₦2,500,000", "—", "₦250,000"],
            ["IPC #1 paid (net)",
             "₦2,125,000",
             "₦2,500,000",
             "₦2,125,000",
             "₦250,000"],
            ["IPC #2 gross", "₦2,500,000", "₦5,000,000", "—", "₦500,000"],
            ["IPC #2 paid (net)",
             "₦2,125,000",
             "₦5,000,000",
             "₦4,250,000",
             "₦500,000"],
            ["Variation #1 approved (ADDITION)", "+₦1,000,000",
             "—", "—", "ceiling ↑ ₦11,000,000"],
            ["Practical Completion", "—", "₦5,000,000", "₦4,250,000", "₦500,000"],
            ["Practical retention release",
             "₦250,000",
             "₦5,000,000",
             "₦4,500,000",
             "₦250,000"],
            ["Final Completion", "—", "—", "—", "—"],
            ["Final retention release",
             "₦250,000",
             "₦5,000,000",
             "₦4,750,000",
             "₦0"],
            ["Contract CLOSED", "—", "₦5,000,000", "₦4,750,000", "₦0"],
        ],
        widths_cm=[5.5, 2.5, 2.7, 2.3, 2.3],
    )

    add_callout(
        doc,
        "Net payable math (IPC #1)",
        "gross = 2,500,000 · retention (10%) = 250,000 · mob recovery (if any) = 0\n"
        "this_certificate_net = 2,500,000 − 250,000 − 0 = 2,250,000\n"
        "WHT at payment (5%) = 125,000 → bank pays contractor 2,125,000; FIRS receives 125,000."
    )

    # ── Section 8: Roles & Responsibilities ──────────────────────────
    add_heading(doc, "8. Roles & Responsibilities (RACI)", level=1)
    add_table(
        doc,
        header=["Activity", "Responsible", "Accountable", "Consulted", "Informed"],
        rows=[
            ["Contract creation", "Procurement Officer", "MDA Head", "Legal; Finance", "Vendor"],
            ["Activation", "Procurement Officer", "MDA Head", "Finance", "Vendor; AG"],
            ["Mobilization payment", "AP Officer", "Accountant-General", "Treasury", "Vendor"],
            ["IPC preparation", "Site Engineer", "Project Engineer", "Contractor", "—"],
            ["IPC certification", "Project Engineer", "MDA Engineer", "Measurement Book", "AP"],
            ["IPC approval", "Finance Approver", "MDA Head / PS", "Internal Audit", "Treasury"],
            ["Voucher raising", "AP Officer", "Accountant-General", "Treasury", "Audit"],
            ["Payment", "Treasury / Cashier", "Accountant-General", "Bank; FIRS", "Vendor"],
            ["Variation approval", "Tier actor", "MDA / Board / BPP", "Engineer", "Vendor; AG"],
            ["Retention release", "AP + Treasury", "Accountant-General", "Engineer", "Vendor"],
            ["Closure", "Procurement Officer", "MDA Head", "Audit", "Vendor; AG"],
        ],
        widths_cm=[3.5, 3.0, 3.0, 3.5, 3.0],
    )

    # ── Section 9: API & Service Reference ──────────────────────────
    doc.add_page_break()
    add_heading(doc, "9. API & Service Reference", level=1)
    add_table(
        doc,
        header=["Step", "HTTP Endpoint", "Service call"],
        rows=[
            ["Create contract", "POST /api/v1/contracts/contracts/", "ContractSerializer.create"],
            ["Activate", "POST /contracts/{id}/activate/", "ContractActivationService.activate"],
            ["Create milestones", "POST /api/v1/contracts/milestones/", "MilestoneSerializer.create"],
            ["Transition state", "POST /contracts/{id}/transition/", "Contract.transition_to"],
            ["Raise mobilization", "POST /contracts/{id}/mobilization/raise/", "MobilizationService.raise_voucher"],
            ["Pay mobilization", "POST /mobilization/{id}/pay/", "MobilizationService.mark_paid"],
            ["Submit IPC", "POST /ipcs/", "IPCService.submit_ipc"],
            ["Certify IPC", "POST /ipcs/{id}/certify/", "IPCService.certify"],
            ["Approve IPC", "POST /ipcs/{id}/approve/", "IPCService.approve"],
            ["Raise voucher", "POST /ipcs/{id}/raise-voucher/", "IPCService.raise_voucher"],
            ["Mark paid", "POST /ipcs/{id}/mark-paid/", "IPCService.mark_paid"],
            ["Submit variation", "POST /variations/", "ContractVariation.save"],
            ["Approve variation", "POST /variations/{id}/approve/", "VariationService.approve"],
            ["Request retention release", "POST /retentions/", "RetentionService.request"],
            ["Pay retention", "POST /retentions/{id}/pay/", "RetentionService.mark_paid"],
            ["Close contract", "POST /contracts/{id}/close/", "Contract.transition_to(CLOSED)"],
        ],
        widths_cm=[3.8, 6.0, 6.2],
    )

    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(
        "Quot PSE · Nigeria IFMIS · Delta State pilot — "
        "process derived directly from contracts/, accounting/, and workflow/ source."
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = GREY

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[OK] Wrote {out_path}")


if __name__ == "__main__":
    here = Path(__file__).resolve().parent.parent
    out = here / "docs" / "Contracts_Milestone_Payment_Process.docx"
    build_document(out)
